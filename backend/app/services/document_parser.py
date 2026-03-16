import io
from pathlib import Path

import openpyxl
from docx import Document
from pypdf import PdfReader


class DocumentParser:
    """Parse various document formats and extract text content."""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from a PDF file using pypdf."""
        reader = PdfReader(file_path)
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages_text.append(f"--- Page {i + 1} ---\n{text}")
        return "\n\n".join(pages_text)

    @staticmethod
    def parse_excel(file_path: str) -> str:
        """Extract text from an Excel file using openpyxl."""
        wb = openpyxl.load_workbook(file_path, data_only=True)
        sheets_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            rows_text.append(f"=== Sheet: {sheet_name} ===")

            for row in ws.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    rows_text.append("\t".join(row_values))

            sheets_text.append("\n".join(rows_text))

        return "\n\n".join(sheets_text)

    @staticmethod
    def parse_word(file_path: str) -> str:
        """Extract text from a Word (.docx) file using python-docx."""
        doc = Document(file_path)
        sections = []

        for para in doc.paragraphs:
            if para.text.strip():
                sections.append(para.text)

        # Also extract tables
        for i, table in enumerate(doc.tables):
            sections.append(f"\n--- Table {i + 1} ---")
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    sections.append(row_text)

        return "\n".join(sections)

    @classmethod
    def parse(cls, file_path: str, content_type: str) -> str:
        """
        Parse a document based on its content type.

        Args:
            file_path: Absolute path to the saved file
            content_type: MIME type of the file

        Returns:
            Extracted text content

        Raises:
            ValueError: If the content type is not supported
        """
        ct = content_type.lower()

        if ct == "application/pdf":
            return cls.parse_pdf(file_path)
        elif ct in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ):
            return cls.parse_excel(file_path)
        elif ct in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return cls.parse_word(file_path)
        else:
            # Fallback: try to infer from file extension
            suffix = Path(file_path).suffix.lower()
            if suffix == ".pdf":
                return cls.parse_pdf(file_path)
            elif suffix in (".xlsx", ".xls"):
                return cls.parse_excel(file_path)
            elif suffix in (".docx", ".doc"):
                return cls.parse_word(file_path)
            else:
                raise ValueError(
                    f"Unsupported document type: {content_type} (extension: {suffix})"
                )


document_parser = DocumentParser()
